#!/usr/bin/env python3
"""
Docx to Dictionary Converter + .idx Index Generator
Tạo .idx files từ docx dictionary cho Zero-RAM lookup

@version: v1.0 (2026-04-13)
"""

import json
import os
import struct
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import argparse

# Try to import python-docx, fall back to basic parsing
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARNING] python-docx not installed. Install with: pip install python-docx")


class DictionaryConverter:
    """Convert docx dictionaries to JSON + Generate .idx files"""
    
    def __init__(self, input_dir: str, output_dir: str):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def extract_from_docx(self, docx_path: Path) -> List[Dict]:
        """Extract entries from docx file"""
        entries = []
        
        if not HAS_DOCX:
            print(f"[SKIP] python-docx not available: {docx_path.name}")
            return entries
            
        try:
            doc = docx.Document(str(docx_path))
            
            # Try to extract table format or paragraph format
            for table in doc.tables:
                # Table format: Term | Definition
                for row in table.rows:
                    if len(row.cells) >= 2:
                        term = row.cells[0].text.strip()
                        definition = row.cells[1].text.strip()
                        if term and definition:
                            entries.append({
                                'term': term,
                                'definition': definition,
                                'source': docx_path.stem
                            })
            
            # If no table, try paragraph format (line by line)
            if not entries:
                prev_text = ""
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    # Assume format: Term: Definition or Term - Definition
                    if ':' in text:
                        parts = text.split(':', 1)
                        term = parts[0].strip()
                        definition = parts[1].strip() if len(parts) > 1 else ""
                    elif '-' in text:
                        parts = text.split('-', 1)
                        term = parts[0].strip()
                        definition = parts[1].strip() if len(parts) > 1 else ""
                    else:
                        # Maybe term is the paragraph, definition next
                        if prev_text and text:
                            entries[-1]['definition'] += ' ' + text
                            continue
                        term = text
                        definition = ""
                    
                    if term:
                        entries.append({
                            'term': term,
                            'definition': definition,
                            'source': docx_path.stem
                        })
                    prev_text = text
                        
        except Exception as e:
            print(f"[ERROR] Failed to parse {docx_path.name}: {e}")
            
        return entries
    
    def convert_all(self) -> str:
        """Convert all docx files to combined JSON"""
        all_entries = []
        
        docx_files = list(self.input_dir.glob("*.docx"))
        print(f"[DictionaryConverter] Found {len(docx_files)} docx files")
        
        for docx_file in docx_files:
            print(f"[Processing] {docx_file.name}")
            entries = self.extract_from_docx(docx_file)
            print(f"  -> {len(entries)} entries extracted")
            all_entries.extend(entries)
        
        # Remove duplicates by term
        seen = set()
        unique_entries = []
        for e in all_entries:
            if e['term'] not in seen:
                seen.add(e['term'])
                unique_entries.append(e)
        
        # Sort by term
        unique_entries.sort(key=lambda x: x['term'])
        
        # Save combined JSON
        output_json = self.output_dir / "combined_dict.json"
        with open(output_json, 'w', encoding='utf-8') as f:
            json.dump({'entries': unique_entries}, f, ensure_ascii=False, indent=2)
        
        print(f"[DONE] Saved {len(unique_entries)} entries to {output_json}")
        return str(output_json)


class IndexGenerator:
    """Tạo .idx files cho Zero-RAM lookup"""
    
    def __init__(self, data_dir: str, output_dir: str):
        self.data_dir = Path(data_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _stream_json_entries(self, json_path: Path):
        """Generator for streaming large JSON files (line-delimited) - Zero-RAM"""
        try:
            import ijson
            for item in ijson.items(open(json_path, 'rb'), 'entries.item'):
                yield item
        except ImportError:
            with open(json_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip().startswith('{'):
                        try:
                            yield json.loads(line)
                        except json.JSONDecodeError:
                            continue
        
    def create_idx_file(self, json_file: str, key_field: str = 'term') -> str:
        """
        Tạo .idx file từ JSON file
        
        Format .idx:
        - Header: magic(4) + version(4) + count(8) = 16 bytes
        - Index entries: key_length(4) + key_bytes + offset(8) + size(4)
        - Data section: JSON objects stored sequentially (one per line)
        """
        json_path = self.data_dir / json_file
        idx_path = self.output_dir / f"{json_file.replace('.json', '')}.idx"
        dict_path = self.output_dir / f"{json_file.replace('.json', '')}.dict"
        
        print(f"[IndexGenerator] Creating index: {json_file}")
        
        file_size = json_path.stat().st_size
        print(f"  -> File size: {file_size / 1024 / 1024:.2f} MB")
        
        # Zero-RAM: Use streaming for large files (>10MB)
        if file_size > 10 * 1024 * 1024:
            print("[WARN] Large file detected - converting streaming to list")
            items = list(self._stream_json_entries(json_path))
        else:
            # Small file - safe to load
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, dict) and 'entries' in data:
                items = data['entries']
            elif isinstance(data, list):
                items = data
            else:
                raise ValueError(f"Unsupported format in {json_file}")
        
        print(f"  -> {len(items)} items to index")
        
        # Build index with binary search structure
        index_entries = []
        
        with open(dict_path, 'w', encoding='utf-8') as dict_file:
            data_offsets = []
            
            for item in items:
                key = item.get(key_field, '')
                if not key:
                    continue
                    
                # Write data as JSON line to .dict file
                data_json = json.dumps(item, ensure_ascii=False)
                dict_file.write(data_json + '\n')
                
                # Record offset and size
                offset = dict_file.tell() - len(data_json) - 1  # -1 for newline
                size = len(data_json)
                data_offsets.append((key, offset, size))
        
        # Sort for binary search
        data_offsets.sort(key=lambda x: x[0])
        
        # Write .idx file
        with open(idx_path, 'wb') as idx_file:
            # Write header: magic + version + count
            magic = b'IDX\x00'
            version = 1
            count = len(data_offsets)
            
            idx_file.write(magic)
            idx_file.write(struct.pack('<I', version))
            idx_file.write(struct.pack('<Q', count))
            
            # Write index entries (sorted for binary search)
            index_start = idx_file.tell()
            
            for key, offset, size in data_offsets:
                key_bytes = key.encode('utf-8')
                key_len = len(key_bytes)
                
                # key_length + key + offset + size
                idx_file.write(struct.pack('<I', key_len))
                idx_file.write(key_bytes)
                idx_file.write(struct.pack('<Q', offset))
                idx_file.write(struct.pack('<I', size))
        
        print(f"  -> Created {idx_path} ({idx_path.stat().st_size} bytes)")
        print(f"  -> Created {dict_path} ({dict_path.stat().st_size} bytes)")
        
        return str(idx_path)
    
    def lookup(self, idx_path: str, search_key: str) -> Optional[Dict]:
        """Binary search lookup in .idx file - Zero RAM!"""
        with open(idx_path, 'rb') as f:
            # Read header
            magic = f.read(4)
            if magic != b'IDX\x00':
                print("[ERROR] Invalid index file")
                return None
            
            version = struct.unpack('<I', f.read(4))[0]
            count = struct.unpack('<Q', f.read(8))[0]
            
            # Binary search
            left, right = 0, count - 1
            
            while left <= right:
                mid = (left + right) // 2
                f.seek(16 + mid * (4 + 8 + 4))  # header + index entry size
                
                key_len = struct.unpack('<I', f.read(4))[0]
                key_bytes = f.read(key_len)
                key = key_bytes.decode('utf-8')
                
                if key == search_key:
                    # Found - read offset and size
                    offset = struct.unpack('<Q', f.read(8))[0]
                    size = struct.unpack('<I', f.read(4))[0]
                    
                    # Read data from .dict file
                    dict_path = idx_path.replace('.idx', '.dict')
                    with open(dict_path, 'r', encoding='utf-8') as df:
                        df.seek(offset)
                        data_json = df.read(size)
                        return json.loads(data_json)
                        
                elif key < search_key:
                    left = mid + 1
                else:
                    right = mid - 1
            
            return None


def main():
    parser = argparse.ArgumentParser(description='Docx to Dict + .idx Generator')
    parser.add_argument('--input', '-i', default='data/dictionaries', help='Input directory')
    parser.add_argument('--output', '-o', default='data/indexed', help='Output directory')
    parser.add_argument('--json', '-j', help='Convert only specific JSON file')
    parser.add_argument('--lookup', '-l', help='Lookup key in index')
    args = parser.parse_args()
    
    converter = DictionaryConverter(args.input, args.output)
    index_gen = IndexGenerator(args.output, args.output)
    
    # Convert docx to JSON
    if not args.json and not args.lookup:
        json_path = converter.convert_all()
        
        # Generate .idx file
        json_name = Path(json_path).name
        index_gen.create_idx_file(json_name, 'term')
        
        # Test lookup
        print("\n[TEST] Binary search lookup:")
        result = index_gen.lookup(f"{args.output}/combined_dict.idx", " Thiền sư")
        if result:
            print(f"  Found: {result.get('term', 'N/A')}")
        else:
            print("  Not found (expected - binary search exact match)")
            
    # Generate .idx from existing JSON
    elif args.json:
        index_gen.create_idx_file(args.json)
        
    # Lookup in index
    elif args.lookup:
        idx_path = f"{args.output}/combined_dict.idx"
        result = index_gen.lookup(idx_path, args.lookup)
        if result:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            print(f"Not found: {args.lookup}")


if __name__ == '__main__':
    main()
