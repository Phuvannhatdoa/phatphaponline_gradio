#!/usr/bin/env python3
"""
D1: Convert .docx → .txt
Convert 24 dictionary files from .docx to .txt for StarDict processing
"""

import os
import sys
import re

try:
    import docx
except ImportError:
    print("📦 Installing python-docx...")
    os.system("pip install python-docx -q")
    import docx

# Paths
SOURCE_DIR = "/opt/phatphaponline_gradio/truyenthua/visjs-app/Dai_Tang_Kinh/daoanh/data/dictionaries"
OUTPUT_DIR = "/opt/phatphaponline_gradio/truyenthua/visjs-app/Dai_Tang_Kinh/daoanh/data/raw/dictionaries"

def extract_text_from_docx(doc_path):
    """Extract text from .docx file"""
    try:
        doc = docx.Document(doc_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"  ⚠️ Error reading {doc_path}: {e}")
        return None

def clean_text(text):
    """Clean and normalize text"""
    lines = text.split('\n')
    cleaned = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip empty lines and very short lines
        if len(line) < 2:
            continue
        
        cleaned.append(line)
    
    return '\n'.join(cleaned)

def convert_all():
    """Convert all .docx files in source directory"""
    print("🚀 D1: Convert .docx → .txt")
    print("=" * 50)
    
    # Create output directory if not exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Get all .docx files
    docx_files = [f for f in os.listdir(SOURCE_DIR) if f.endswith('.docx')]
    
    print(f"📂 Found {len(docx_files)} .docx files")
    
    converted = 0
    failed = 0
    
    for filename in docx_files:
        source_path = os.path.join(SOURCE_DIR, filename)
        output_filename = filename.replace('.docx', '.txt')
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        print(f"  🔄 Converting: {filename}")
        
        text = extract_text_from_docx(source_path)
        
        if text:
            # Clean text
            text = clean_text(text)
            
            # Write to output
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"     ✅ → {output_filename} ({len(text)} chars)")
            converted += 1
        else:
            print(f"     ❌ Failed")
            failed += 1
    
    print(f"\n✅ Complete!")
    print(f"   Converted: {converted}")
    print(f"   Failed: {failed}")
    print(f"   Output: {OUTPUT_DIR}")
    
    return converted, failed

if __name__ == "__main__":
    convert_all()