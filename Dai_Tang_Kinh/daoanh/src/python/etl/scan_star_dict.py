#!/usr/bin/env python3
"""
T1: Scan StarDict Files - Quét 23 file .docx, trích xuất địa danh Chùa/Tự/Viện
Input:  data/dictionaries/*.docx
Output: data/processed/raw_temples.json

Quy trình:
1. Đọc danh sách file .docx trong thư mục dictionaries
2. Parse từng file bằng python-docx
3. Áp dụng bộ lọc kép:
   - Điều kiện 1 (Tên): Tiền tố/Hậu tố địa danh
   - Điều kiện 2 (Ngữ cảnh): Từ khóa địa lý
4. Trích xuất metadata (tên, mô tả, nguồn)
"""

import os
import re
import json
import glob
from pathlib import Path
from datetime import datetime
import sys

# Cấu hình đường dẫn
BASE_DIR = Path("/opt/phatphaponline_gradio/truyenthua/visjs-app/Dai_Tang_Kinh/daoanh")
DICT_DIR = BASE_DIR / "data" / "dictionaries"
OUTPUT_DIR = BASE_DIR / "data" / "processed"
LOG_DIR = BASE_DIR / "logs"

# Đảm bảo thư mục tồn tại
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

# ============================================
# BỘ LỌC KÉP (DUAL FILTER)
# ============================================

# Điều kiện 1: Tiền tố và hậu tố của tên địa danh
PREFIXES = [
    "Chùa", "Tịnh Xá", "Thiền Viện", "Tổ Đình", "Am", "Cốc", 
    "Đại Tự", "Thánh Địa", "Thắng Phúc", "Phước Định",
    "Viện", "Quán", "Trai", "Xá", "Tinh Xá", "Nirvana"
]

SUFFIXES = [
    "Tự", "Viện", "Quán", "Trai", "Xá", "Tinh Xá", 
    "Am", "Cốc", "Động", "Quả", "Nham", "Sơn"
]

# Điều kiện 2: Từ khóa ngữ cảnh địa lý (phải có ít nhất 1)
CONTEXT_KEYWORDS = [
    "tọa lạc", "ở tại", "tọa độ", "nằm tại", "vị trí",
    "thuộc tỉnh", "thuộc huyện", "thuộc phường", "thuộc xã",
    "xây dựng", "kiến trúc", "trùng tu", "tôn tạo",
    "núi", "rừng", "sông", "biển", "đồi", "cao",
    "xã", "thôn", "phường", "quận", "huyện", "tỉnh",
    "thành phố", "tp.", "tp ", "tp."
]

# ============================================
# HÀM HỖ TRỢ
# ============================================

def remove_diacritics(text):
    """Loại bỏ dấu tiếng Việt để so khớp"""
    # Mapping dấu tiếng Việt
    import unicodedata
    nfd = unicodedata.normalize('NFD', text)
    return ''.join(ch for ch in nfd if unicodedata.category(ch) != 'Mn')

def is_temple_name(name: str) -> bool:
    """Kiểm tra xem tên có phải địa danh chùa không (Điều kiện 1)"""
    name = name.strip()
    
    # Kiểm tra tiền tố
    for prefix in PREFIXES:
        if name.startswith(prefix):
            return True
    
    # Kiểm tra hậu tố
    for suffix in SUFFIXES:
        if name.endswith(suffix):
            return True
    
    return False

def has_location_context(text: str) -> bool:
    """Kiểm tra xem text có ngữ cảnh địa lý không (Điều kiện 2)"""
    text_lower = text.lower()
    for keyword in CONTEXT_KEYWORDS:
        if keyword.lower() in text_lower:
            return True
    return False

def extract_han_viet(name: str) -> str:
    """Trích xuất tên Hán-Việt từ trong ngoặc đơn nếu có"""
    match = re.search(r'\(([^)]+)\)', name)
    if match:
        return match.group(1).strip()
    return name.strip()

def clean_temple_name(name: str) -> str:
    """Làm sạch tên chùa - loại bỏ ký tự đặc biệt"""
    # Loại bỏ các ký tự đặc biệt nhưng giữ tiếng Việt
    name = re.sub(r'[^\w\sÀ-ỹ]', '', name)
    name = ' '.join(name.split())  # Loại bỏ khoảng trắng thừa
    return name.strip()

# ============================================
# XỬ LÝ FILE DOCX
# ============================================

def read_docx_file(filepath: Path) -> list:
    """Đọc file .docx và trả về danh sách các mục từ điển"""
    try:
        import docx
    except ImportError:
        print("⚠️ Chưa cài python-docx. Đang cài đặt...")
        os.system(f"{sys.executable} -m pip install python-docx -q")
        import docx
    
    items = []
    try:
        doc = docx.Document(str(filepath))
        
        # Đọc từng đoạn văn
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                items.append(text)
        
        # Đọc từng bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join([cell.text.strip() for cell in row.cells])
                if row_text:
                    items.append(row_text)
    
    except Exception as e:
        print(f"  ❌ Lỗi đọc file {filepath.name}: {e}")
    
    return items

def extract_temples_from_items(items: list, source_file: str) -> list:
    """Trích xuất địa danh từ danh sách items theo bộ lọc kép"""
    temples = []
    
    for item in items:
        # Tách keyword và value (tab hoặc dấu hai chấm)
        parts = re.split(r'[\t:]+', item, maxsplit=1)
        
        if len(parts) == 2:
            keyword = parts[0].strip()
            value = parts[1].strip()
        else:
            keyword = item
            value = item
        
        # Áp dụng bộ lọc kép
        if is_temple_name(keyword) and has_location_context(value):
            temple = {
                "name": clean_temple_name(keyword),
                "han_viet": extract_han_viet(keyword),
                "description": value.strip()[:500],  # Giới hạn 500 ký tự
                "source": source_file,
                "raw_keyword": keyword,
                "raw_value": value
            }
            temples.append(temple)
    
    return temples

# ============================================
# MAIN PROCESS
# ============================================

def scan_all_dictionaries():
    """Quét tất cả file từ điển và trích xuất địa danh"""
    print("=" * 60)
    print("🔍 T1: QUÉT FILE TỪ ĐIỂN - TRÍCH XUẤT ĐỊA DANH")
    print("=" * 60)
    print(f"📁 Thư mục: {DICT_DIR}")
    print()
    
    # Lấy danh sách file
    docx_files = sorted(DICT_DIR.glob("*.docx"))
    
    if not docx_files:
        print("❌ Không tìm thấy file .docx nào!")
        return None
    
    print(f"📊 Tìm thấy {len(docx_files)} file .docx")
    print("-" * 60)
    
    all_temples = []
    file_stats = []
    
    for idx, filepath in enumerate(docx_files, 1):
        print(f"  [{idx:2d}/{len(docx_files)}] Đang xử lý: {filepath.name}")
        
        # Đọc file
        items = read_docx_file(filepath)
        
        # Trích xuất địa danh
        temples = extract_temples_from_items(items, filepath.name)
        
        print(f"           → Tìm thấy {len(temples)} địa danh")
        
        file_stats.append({
            "file": filepath.name,
            "items_count": len(items),
            "temples_found": len(temples)
        })
        
        all_temples.extend(temples)
    
    print("-" * 60)
    print(f"✅ Tổng cộng: {len(all_temples)} địa danh được trích xuất")
    print()
    
    # Loại bỏ trùng lặp (dựa trên tên)
    unique_temples = []
    seen_names = set()
    
    for temple in all_temples:
        name_key = remove_diacritics(temple["name"]).lower()
        if name_key not in seen_names:
            seen_names.add(name_key)
            unique_temples.append(temple)
    
    print(f"📊 Sau khi loại trùng: {len(unique_temples)} địa danh duy nhất")
    print()
    
    # Lưu kết quả
    output_data = {
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "source_files": len(docx_files),
            "total_extracted": len(all_temples),
            "unique_temples": len(unique_temples),
            "task": "T1 - Scan StarDict"
        },
        "file_stats": file_stats,
        "temples": unique_temples
    }
    
    output_file = OUTPUT_DIR / "raw_temples.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print(f"💾 Đã lưu: {output_file}")
    print()
    
    # Log
    log_file = LOG_DIR / "task_T1_scan.log"
    with open(log_file, "w", encoding="utf-8") as f:
        f.write(f"Task T1 - Scan StarDict - {datetime.now()}\n")
        f.write(f"Files processed: {len(docx_files)}\n")
        f.write(f"Temples extracted: {len(all_temples)}\n")
        f.write(f"Unique temples: {len(unique_temples)}\n")
        f.write("\nFile Stats:\n")
        for stat in file_stats:
            f.write(f"  {stat['file']}: {stat['temples_found']} temples\n")
    
    print(f"📝 Log: {log_file}")
    
    return output_data

# ============================================
# CHẠY SCRIPT
# ============================================

if __name__ == "__main__":
    try:
        result = scan_all_dictionaries()
        
        if result:
            print("=" * 60)
            print("✅ HOÀN THÀNH T1: Scan StarDict")
            print("=" * 60)
            print(f"📊 Tổng địa danh: {result['metadata']['unique_temples']}")
            print(f"📁 Output: data/processed/raw_temples.json")
        else:
            print("❌ Thất bại!")
            
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()